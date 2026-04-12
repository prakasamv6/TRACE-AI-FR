"""
TRACE-AI-FR — Word Document (.docx) Question Parser

Reads examination questions from a Word document and returns them
as a list of ExaminationQuestion objects.
"""
import re
import logging
from pathlib import Path
from typing import List

from .models import ExaminationQuestion

logger = logging.getLogger(__name__)

# Pattern: lines starting with a number followed by . or ) or :
_QUESTION_RE = re.compile(
    r"^\s*(?:Q(?:uestion)?\s*)?(\d+)\s*[.):\-]\s*(.+)", re.IGNORECASE
)


def parse_questions_from_docx(docx_path: str) -> List[ExaminationQuestion]:
    """
    Parse examination questions from a Word document.

    Supported formats:
      - Numbered lines:  ``1. What browser...``  ``Q2) Was ChatGPT...``
      - Paragraphs that start with a digit + separator
      - Plain paragraphs (treated as individual questions when no numbering)

    Returns a list of ExaminationQuestion with .number and .text populated.
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx is not installed. Cannot parse Word documents.")
        raise RuntimeError(
            "python-docx is required to read Word documents. "
            "Install it with:  pip install python-docx"
        )

    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"Questions document not found: {docx_path}")
    if path.suffix.lower() not in (".docx", ".doc"):
        raise ValueError(f"Unsupported file type: {path.suffix}. Expected .docx")

    doc = Document(str(path))
    questions: List[ExaminationQuestion] = []
    current_text_lines: List[str] = []
    current_num: int = 0
    auto_num = 0

    def _flush():
        nonlocal current_text_lines, current_num, auto_num
        text = " ".join(current_text_lines).strip()
        if text:
            if current_num == 0:
                auto_num += 1
                current_num = auto_num
            else:
                auto_num = current_num
            questions.append(ExaminationQuestion(number=current_num, text=text))
        current_text_lines = []
        current_num = 0

    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            _flush()
            continue

        m = _QUESTION_RE.match(line)
        if m:
            _flush()
            current_num = int(m.group(1))
            current_text_lines.append(m.group(2).strip())
        else:
            if not current_text_lines:
                # Start new question from plain paragraph
                _flush()
            current_text_lines.append(line)

    _flush()  # last accumulated question

    # Also extract from tables (some question docs use tables)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if not cells:
                continue
            # Try to find numbered question in first cell
            combined = " ".join(cells)
            m = _QUESTION_RE.match(combined)
            if m:
                auto_num += 1
                questions.append(ExaminationQuestion(
                    number=int(m.group(1)),
                    text=m.group(2).strip(),
                ))
            elif len(combined) > 10:  # skip headers / short labels
                auto_num += 1
                questions.append(ExaminationQuestion(
                    number=auto_num,
                    text=combined,
                ))

    logger.info(f"Parsed {len(questions)} questions from {path.name}")
    return questions
