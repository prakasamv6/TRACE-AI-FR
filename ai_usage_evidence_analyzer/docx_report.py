"""
Word (.docx) forensic report generator.

Produces a court-ready forensic examination report in the UNCC Digital
Forensics Laboratory template format (Cambria, centered header, section
headings bold/underlined, numbered figures, appendix with glossary).

Formatted for submission to U.S. District Courts, state courts, and
legal counsel in compliance with Federal Rules of Evidence (FRE 702)
and Daubert/Kumho standards.
"""

from __future__ import annotations

import logging
import os
from datetime import datetime
from typing import List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .models import ForensicReport

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants — Court-standard formatting
# ---------------------------------------------------------------------------
_FONT = "Cambria"
_BODY_SIZE = Pt(12)
_HEADING_SIZE = Pt(13)
_HEADER_SIZE = Pt(14)
_TITLE_SIZE = Pt(16)
_SMALL_SIZE = Pt(10)
_TABLE_SIZE = Pt(9)
_CONFIDENTIAL_COLOR = RGBColor(0xC0, 0x00, 0x00)  # Dark red
_HEADING_COLOR = RGBColor(0x1A, 0x1A, 0x5C)  # Navy
_BORDER_COLOR = RGBColor(0x4F, 0x81, 0xBD)  # Professional blue

_SECTION_COUNTER = 0  # Module-level counter for auto-numbering sections


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def _reset_section_counter():
    global _SECTION_COUNTER
    _SECTION_COUNTER = 0


def _next_section_number() -> str:
    global _SECTION_COUNTER
    _SECTION_COUNTER += 1
    return str(_SECTION_COUNTER)


def _set_run(run, *, bold=False, italic=False, underline=False,
             size=_BODY_SIZE, font=_FONT, color=None):
    """Apply formatting to a single run."""
    run.font.name = font
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = color


def _add_centered(doc, text, *, bold=False, underline=False,
                   size=_BODY_SIZE, spacing_after=Pt(4), color=None):
    """Add a centered paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = spacing_after
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    _set_run(run, bold=bold, underline=underline, size=size, color=color)
    return p


def _add_heading_styled(doc, text, *, spacing_before=Pt(18), numbered=True):
    """Add a numbered section heading (bold, underlined, Cambria 13pt, navy)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = spacing_before
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    prefix = f"{_next_section_number()}. " if numbered else ""
    run = p.add_run(f"{prefix}{text}")
    _set_run(run, bold=True, underline=True, size=_HEADING_SIZE,
             color=_HEADING_COLOR)
    return p


def _add_subheading(doc, text, *, spacing_before=Pt(10)):
    """Add a sub-section heading (bold, Cambria 12pt)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = spacing_before
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    _set_run(run, bold=True, size=_BODY_SIZE, color=_HEADING_COLOR)
    return p


def _add_body(doc, text, *, spacing_after=Pt(6), indent=None):
    """Add a normal body paragraph with optional indentation."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = spacing_after
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = indent
    run = p.add_run(text)
    _set_run(run, size=_BODY_SIZE)
    return p


def _add_bullet(doc, text, *, spacing_after=Pt(3), level=0):
    """Add a bullet-point body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = spacing_after
    p.paragraph_format.space_before = Pt(0)
    indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.left_indent = indent
    p.paragraph_format.first_line_indent = Inches(-0.2)
    run = p.add_run(f"\u2022  {text}")
    _set_run(run, size=_BODY_SIZE)
    return p


def _add_blank(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def _add_horizontal_line(doc):
    """Add a thin horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="4F81BD"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _style_table_header(cell):
    """Style a table header cell with navy background and white text."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="1A1A5C" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = _FONT
            run.font.size = _TABLE_SIZE
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _style_table_cell(cell):
    """Style a regular table cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = _FONT
            run.font.size = _TABLE_SIZE


def _add_page_break(doc):
    """Insert a page break."""
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_docx_report(
    report: ForensicReport,
    output_path: str,
    examiner_name: str = "",
    examination_name: str = "",
    in_the_matter_of: str = "",
    organization: str = "University of North Carolina at Charlotte",
    lab_name: str = "Digital Forensics Laboratory",
    lab_address: str = "9201 University Blvd, Charlotte NC 28223",
) -> str:
    """
    Generate a Word forensic report following the UNCC DFL template.

    The language is written for non-technical readers such as judges,
    attorneys, and jurors. Technical terms are explained in plain English
    and a glossary is provided in the appendix.

    Returns the path to the generated .docx file.
    """
    _reset_section_counter()
    doc = Document()

    # ── Page setup ─────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(0.75)

    # ── Header ─────────────────────────────────────────────────────
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()

    h1 = header.paragraphs[0]
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h1.add_run(organization)
    _set_run(run, size=_BODY_SIZE, bold=True)

    h2 = header.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h2.add_run(lab_name)
    _set_run(run, size=_SMALL_SIZE)

    h3 = header.add_paragraph()
    h3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h3.add_run(lab_address)
    _set_run(run, size=_SMALL_SIZE)

    # ── Footer ─────────────────────────────────────────────────────
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    matter_text = in_the_matter_of or "(Subject)"
    run = fp.add_run(f"CONFIDENTIAL — In the matter of: {matter_text}")
    _set_run(run, size=_SMALL_SIZE, italic=True, color=_CONFIDENTIAL_COLOR)

    # ── Title page ─────────────────────────────────────────────────
    _add_blank(doc, 2)
    _add_centered(doc, "FORENSIC EXAMINATION REPORT",
                  bold=True, size=_TITLE_SIZE, color=_HEADING_COLOR)
    _add_horizontal_line(doc)
    _add_blank(doc)

    _add_centered(doc, "In the matter of", bold=True, underline=True,
                  size=_HEADER_SIZE)
    _add_centered(doc, in_the_matter_of or "(Subject)",
                  bold=True, size=_HEADER_SIZE)
    _add_blank(doc)

    examiner_display = examiner_name or report.case_info.examiner or "Digital Forensics Examiner"
    _add_centered(doc, "Prepared by:", size=_BODY_SIZE)
    _add_centered(doc, examiner_display, bold=True, size=_HEADER_SIZE)
    _add_centered(doc, "Digital Forensics Examiner", size=_BODY_SIZE)
    _add_centered(doc, organization, size=_BODY_SIZE)
    _add_blank(doc)

    report_date = report.generated_at.strftime('%B %d, %Y') if report.generated_at else 'N/A'
    case_id = report.case_info.case_id or "N/A"
    _add_centered(doc, f"Report Date: {report_date}", size=_BODY_SIZE)
    _add_centered(doc, f"Case Reference: {case_id}", size=_BODY_SIZE)

    _add_blank(doc, 2)
    _add_horizontal_line(doc)
    _add_centered(doc, "CONFIDENTIAL — FOR LEGAL USE ONLY",
                  bold=True, size=_BODY_SIZE, color=_CONFIDENTIAL_COLOR)
    _add_horizontal_line(doc)

    # ── Page break before body ────────────────────────────────────
    _add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 1. PURPOSE AND SCOPE
    # ══════════════════════════════════════════════════════════════
    _add_heading_styled(doc, "Purpose and Scope of This Report")

    case_name = examination_name or report.case_info.case_name or "this case"

    _add_body(doc, (
        f"This report presents the findings of a forensic examination "
        f"conducted in connection with {case_name}. The purpose of this "
        f"examination was to determine whether any artificial intelligence "
        f"(AI) services — such as ChatGPT, Claude, or Google Gemini — "
        f"were used on the computer or device that was submitted as evidence."
    ))

    _add_body(doc, (
        "In plain terms, the examiner looked through the digital evidence "
        "for signs that someone visited or used AI websites and applications. "
        "This report explains what was found, how confident the examiner is "
        "in those findings, and what limitations apply."
    ))

    _add_body(doc, (
        f"The examination was performed using a forensic analysis tool called "
        f"TRACE-AI-FR (version {report.tool_version}). This tool reads "
        f"digital evidence in a read-only manner, meaning it does not alter "
        f"or modify the original evidence in any way."
    ))

    # ══════════════════════════════════════════════════════════════
    # 2. SUMMARY OF FINDINGS
    # ══════════════════════════════════════════════════════════════
    _add_heading_styled(doc, "Summary of Findings")

    n_artifacts = len(report.all_artifacts)
    n_fraues = len(report.fraues)
    n_platforms = len(report.ai_footprints)

    platforms_detected = []
    for fprt in report.ai_footprints:
        pname = fprt.platform.value if hasattr(fprt.platform, "value") else str(fprt.platform)
        platforms_detected.append(pname)

    if platforms_detected:
        _add_body(doc, (
            f"The examination found evidence that {n_platforms} AI "
            f"service(s) were accessed from the device under examination. "
            f"The AI services identified are: {', '.join(platforms_detected)}."
        ))
        _add_body(doc, (
            f"In total, {n_artifacts} individual piece(s) of digital evidence "
            f"(called \"artifacts\") were recovered. These artifacts were "
            f"grouped into {n_fraues} distinct finding(s), each representing "
            f"a separate instance or session of AI service usage."
        ))
    else:
        _add_body(doc, (
            f"The examination did not find direct evidence that any AI "
            f"service (such as ChatGPT, Claude, or Google Gemini) was used "
            f"on the device under examination. A total of {n_artifacts} "
            f"digital artifact(s) were reviewed."
        ))
        _add_body(doc, (
            "Important: The absence of evidence does not prove that AI services "
            "were never used. Evidence may have been deleted, the device may "
            "have been cleaned, or AI services may have been accessed through "
            "methods that leave no trace on this device."
        ))

    # ══════════════════════════════════════════════════════════════
    # 3. EVIDENCE EXAMINED
    # ══════════════════════════════════════════════════════════════
    _add_heading_styled(doc, "Evidence Examined")

    ev_path = report.evidence_info.image_path or "(evidence path)"
    ev_name = os.path.basename(ev_path) if ev_path else "N/A"
    ev_size = report.evidence_info.image_size_bytes or 0
    ev_size_str = _human_size(ev_size) if ev_size else "N/A"

    _add_body(doc, (
        f"The examiner received a digital copy of the evidence in the form "
        f"of a forensic image file named \"{ev_name}\" ({ev_size_str}). "
        f"A forensic image is an exact, bit-for-bit copy of a computer's "
        f"storage, preserving all data — including deleted files — exactly "
        f"as they existed on the original device."
    ))

    # EWF metadata
    ewf_meta = getattr(report.evidence_info, "ewf_metadata", {})
    if ewf_meta:
        _add_subheading(doc, "Evidence Details")
        for k, v in ewf_meta.items():
            _add_bullet(doc, f"{k}: {v}")

    # Hash information
    hashes = getattr(report.evidence_info, "hash_values", {})
    if hashes:
        _add_subheading(doc, "Evidence Integrity Verification")
        _add_body(doc, (
            "To confirm that the evidence was not altered during the "
            "examination, the examiner calculated digital fingerprints "
            "(called \"hash values\") of the evidence file. If these "
            "fingerprints match the originals, it proves the evidence "
            "is authentic and unmodified."
        ))
        for algo, val in hashes.items():
            _add_bullet(doc, f"{algo.upper()}: {val}")

    # OS detection
    os_info = getattr(report.evidence_info, "detected_os", None)
    if os_info:
        os_str = os_info.value if hasattr(os_info, "value") else str(os_info)
        _add_body(doc, (
            f"The device was running the {os_str} operating system."
        ))

    # User profiles
    profiles = getattr(report.evidence_info, "user_profiles", [])
    if profiles:
        _add_body(doc, (
            f"The following user account(s) were found on the device: "
            f"{', '.join(profiles)}."
        ))

    # Partitions
    partitions = getattr(report.evidence_info, "partitions", [])
    if partitions:
        _add_body(doc, (
            f"The device's storage was organized into {len(partitions)} "
            f"section(s) (called \"partitions\"):"
        ))
        for pi in partitions:
            desc = getattr(pi, "description", "")
            size = getattr(pi, "length", 0)
            fs = getattr(pi, "filesystem", "")
            label = desc
            if size:
                label += f" — {_human_size(size)}"
            if fs:
                label += f" (file system: {fs})"
            _add_bullet(doc, label)

    # ══════════════════════════════════════════════════════════════
    # 4. EXAMINATION PROCESS
    # ══════════════════════════════════════════════════════════════
    _add_heading_styled(doc, "Examination Process")

    _add_body(doc, (
        "The examiner used specialized software to systematically search "
        "the evidence for any traces of AI service usage. The software "
        "examined the following types of evidence:"
    ))

    _add_bullet(doc, (
        "Web browsing history — Records of websites visited, including "
        "AI service websites such as chat.openai.com (ChatGPT), "
        "claude.ai (Claude), and gemini.google.com (Google Gemini)."
    ))
    _add_bullet(doc, (
        "Internet cookies — Small data files that websites store on a "
        "computer, which can show that a user visited or logged into "
        "an AI service."
    ))
    _add_bullet(doc, (
        "Downloaded files — Files saved from AI services to the "
        "computer, such as conversation exports or AI-generated content."
    ))
    _add_bullet(doc, (
        "Application data — Settings, cached data, and logs left "
        "behind by AI applications installed on the device."
    ))
    _add_bullet(doc, (
        "Recent file activity — Records showing which files were "
        "recently opened, which may include AI-generated documents."
    ))
    _add_bullet(doc, (
        "Deleted items — Files or records that were deleted from the "
        "device but may still be recoverable from the storage."
    ))

    # Parsers executed
    if report.parser_results:
        _add_body(doc, (
            f"The software ran {len(report.parser_results)} different "
            f"analysis module(s), each designed to look for a specific "
            f"type of evidence:"
        ))
        for pr in report.parser_results:
            parser_name = pr.parser_name if hasattr(pr, "parser_name") else str(pr)
            status = ""
            if hasattr(pr, "status"):
                status_val = pr.status.value if hasattr(pr.status, 'value') else pr.status
                status = f" — {status_val}"
            n = ""
            if hasattr(pr, "artifacts_found"):
                n = f" ({pr.artifacts_found} item(s) found)"
            _add_bullet(doc, f"{parser_name}{status}{n}")

    # ══════════════════════════════════════════════════════════════
    # 5. DETAILED FINDINGS
    # ══════════════════════════════════════════════════════════════
    _add_heading_styled(doc, "Detailed Findings")

    # Platform findings table
    if report.ai_footprints:
        _add_subheading(doc, "AI Services Detected")
        _add_body(doc, (
            "The following table summarizes which AI services were found "
            "in the evidence and how confident the examiner is in each finding. "
            "\"Direct\" evidence means the examiner found clear, firsthand "
            "proof (such as a website visit record). \"Supporting\" evidence "
            "means additional clues that strengthen the finding but are not "
            "proof on their own (such as a related cookie)."
        ))

        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        headers = ["AI Service", "Confidence Level",
                    "Total Evidence Items", "Direct Evidence", "Supporting Evidence"]
        for i, h in enumerate(headers):
            hdr[i].text = h
            _style_table_header(hdr[i])

        for fprt in report.ai_footprints:
            row = table.add_row().cells
            pname = fprt.platform.value if hasattr(fprt.platform, "value") else str(fprt.platform)
            conf = fprt.overall_confidence.value if hasattr(fprt.overall_confidence, "value") else str(fprt.overall_confidence)
            row[0].text = pname
            row[1].text = conf
            row[2].text = str(fprt.total_artifacts)
            row[3].text = str(fprt.direct_artifacts)
            row[4].text = str(fprt.inferred_artifacts)
            for cell in row:
                _style_table_cell(cell)

        _add_body(doc, "")

    # FRAUE findings
    if report.fraues:
        _add_subheading(doc, "Individual Findings (Evidence Units)")
        _add_body(doc, (
            "Each row in the table below represents a distinct finding of "
            "AI service usage. The \"Claim Level\" indicates the strength "
            "of the examiner's conclusion — ranging from \"Observed\" "
            "(clearly seen in the evidence) to \"Inferred\" (a reasonable "
            "conclusion based on available clues)."
        ))

        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        headers = ["Finding ID", "AI Service", "Type of Activity",
                    "Date/Time Window", "Confidence", "Claim Strength"]
        for i, h in enumerate(headers):
            hdr[i].text = h
            _style_table_header(hdr[i])

        for f in report.fraues:
            row = table.add_row().cells
            row[0].text = f.fraue_id[:16] + "..."
            row[1].text = f.platform.value if hasattr(f.platform, "value") else str(f.platform)
            row[2].text = f.likely_activity_class or ""
            w_start = f.window_start.strftime("%m/%d/%Y %I:%M %p") if f.window_start else "Unknown"
            w_end = f.window_end.strftime("%m/%d/%Y %I:%M %p") if f.window_end else "Unknown"
            row[3].text = f"{w_start} to {w_end}"
            row[4].text = f.event_confidence.value if hasattr(f.event_confidence, "value") else str(f.event_confidence)
            row[5].text = f.claim_level.value if hasattr(f.claim_level, "value") else str(f.claim_level)
            for cell in row:
                _style_table_cell(cell)

        _add_body(doc, "")

    # Timeline summary
    if report.timeline:
        _add_subheading(doc, "Timeline of Activity")
        _add_body(doc, (
            f"The examination identified {len(report.timeline)} event(s) "
            f"arranged in chronological order. The following is a summary "
            f"of the key events:"
        ))
        for ev in sorted(report.timeline,
                         key=lambda e: e.timestamp if e.timestamp else datetime.min)[:20]:
            ts = ev.timestamp.strftime("%B %d, %Y at %I:%M:%S %p") if ev.timestamp else "Date unknown"
            plat = ev.platform.value if hasattr(ev.platform, "value") else str(ev.platform)
            desc = ev.description or ""
            _add_bullet(doc, f"{ts} — {plat}: {desc}")
        if len(report.timeline) > 20:
            _add_body(doc, (
                f"(Note: {len(report.timeline) - 20} additional event(s) "
                f"are documented in the full analysis data files.)"
            ), indent=Inches(0.5))

    # ══════════════════════════════════════════════════════════════
    # 6. EVIDENCE STRENGTH AND CLASSIFICATION
    # ══════════════════════════════════════════════════════════════
    gr = report.governance_record
    if gr:
        has_evidence_sections = (
            getattr(gr, 'direct_evidence_summary', None)
            or getattr(gr, 'corroborating_evidence_summary', None)
            or getattr(gr, 'missing_evidence_summary', None)
        )
        if has_evidence_sections:
            _add_heading_styled(doc, "Evidence Strength and Classification")

            _add_body(doc, (
                "The evidence found in this examination falls into the "
                "following categories:"
            ))

            if gr.direct_evidence_summary:
                _add_subheading(doc, "Direct Evidence (Strongest)")
                _add_body(doc, (
                    "Direct evidence consists of items that clearly and "
                    "unmistakably show AI service usage, such as a record "
                    "of a website visit to an AI platform."
                ))
                for s in gr.direct_evidence_summary:
                    _add_bullet(doc, s)

            if gr.corroborating_evidence_summary:
                _add_subheading(doc, "Supporting Evidence")
                _add_body(doc, (
                    "Supporting evidence consists of items that, while not "
                    "proof on their own, lend additional support to the "
                    "findings when considered alongside the direct evidence."
                ))
                for s in gr.corroborating_evidence_summary:
                    _add_bullet(doc, s)

            if gr.missing_evidence_summary:
                _add_subheading(doc, "Evidence Not Found")
                _add_body(doc, (
                    "The following types of evidence were expected but "
                    "not found. Their absence may limit the conclusions "
                    "that can be drawn:"
                ))
                for s in gr.missing_evidence_summary:
                    _add_bullet(doc, s)

    # Confidence class summary
    if report.fraues:
        cc_counts = {}
        for f in report.fraues:
            cc = getattr(f, 'confidence_class', None)
            label = cc.value if hasattr(cc, 'value') else str(cc) if cc else "Unclassified"
            cc_counts[label] = cc_counts.get(label, 0) + 1
        if cc_counts:
            _add_subheading(doc, "Confidence Level Summary")
            _add_body(doc, (
                "The following summarizes how confident the examiner is "
                "in each finding:"
            ))
            for label, count in sorted(cc_counts.items()):
                _add_bullet(doc, f"{label}: {count} finding(s)")

    # ══════════════════════════════════════════════════════════════
    # 7. EXAMINATION QUESTIONS
    # ══════════════════════════════════════════════════════════════
    questions = getattr(report, "examination_questions", [])
    if questions:
        _add_heading_styled(doc, "Examination Questions and Answers")

        _add_body(doc, (
            "The following questions were posed as part of this "
            "examination, along with the examiner's findings:"
        ))

        for q in questions:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(f"Question: {q.text}")
            _set_run(run, bold=True, size=_BODY_SIZE)

            if q.answer:
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.25)
                p2.paragraph_format.space_after = Pt(4)
                run2 = p2.add_run(f"Answer: {q.answer}")
                _set_run(run2, size=_BODY_SIZE)

            if q.evidence_references:
                for ref in q.evidence_references:
                    _add_bullet(doc, f"Supporting evidence: {ref}", level=1)

    # ══════════════════════════════════════════════════════════════
    # 8. LIMITATIONS AND IMPORTANT NOTICES
    # ══════════════════════════════════════════════════════════════
    if gr:
        _add_heading_styled(doc, "Limitations and Important Notices")

        _add_body(doc, (
            "The Court should be aware of the following limitations that "
            "apply to this examination. These are provided in the interest "
            "of full transparency and to assist in the proper weighing of "
            "the evidence."
        ))

        if report.scope_of_conclusion:
            _add_subheading(doc, "Scope of Conclusion")
            _add_body(doc, report.scope_of_conclusion)

        if gr.inference_boundaries:
            _add_subheading(doc, "What This Examination Can and Cannot Prove")
            for ib in gr.inference_boundaries:
                _add_bullet(doc, ib)

        if gr.required_disclosures:
            _add_subheading(doc, "Required Disclosures")
            for d in gr.required_disclosures:
                _add_bullet(doc, d)

        if gr.known_blind_spots:
            _add_subheading(doc, "Known Gaps in Analysis")
            _add_body(doc, (
                "The following are areas where this type of analysis has "
                "known limitations:"
            ))
            for b in gr.known_blind_spots:
                _add_bullet(doc, b)

        # v4.0 governance fields
        if getattr(gr, 'provider_capability_blind_spots', None):
            _add_subheading(doc, "AI Service Provider Limitations")
            _add_body(doc, (
                "Each AI service has different features and data practices "
                "that may affect what evidence is available:"
            ))
            for bs in gr.provider_capability_blind_spots:
                _add_bullet(doc, bs)

        if getattr(gr, 'acquisition_blind_spots', None):
            _add_subheading(doc, "Evidence Collection Gaps")
            _add_body(doc, (
                "Certain types of evidence were not available for this "
                "examination:"
            ))
            for bs in gr.acquisition_blind_spots:
                _add_bullet(doc, bs)

        if getattr(gr, 'surface_coverage_summary', None):
            _add_subheading(doc, "Areas of the Device Examined")
            for sc in gr.surface_coverage_summary:
                _add_bullet(doc, sc)

        if getattr(gr, 'alternative_explanations', None):
            _add_subheading(doc, "Alternative Explanations")
            _add_body(doc, (
                "In the interest of fairness, the following alternative "
                "explanations could account for some of the evidence found. "
                "These should be considered when evaluating the findings:"
            ))
            for ae in gr.alternative_explanations:
                _add_bullet(doc, ae)

    # ══════════════════════════════════════════════════════════════
    # 9. VOICE / AUDIO EVIDENCE (if any)
    # ══════════════════════════════════════════════════════════════
    voice_records = getattr(report, 'voice_evidence', [])
    if voice_records:
        _add_heading_styled(doc, "Voice and Audio Evidence")
        _add_body(doc, (
            f"The examination identified {len(voice_records)} audio "
            f"recording(s) that may be related to AI service usage "
            f"(for example, voice conversations with an AI assistant):"
        ))

        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Type", "AI Service", "Source File",
                                "Duration", "Confidence"]):
            hdr[i].text = h
            _style_table_header(hdr[i])
        for vrec in voice_records:
            row = table.add_row().cells
            row[0].text = str(getattr(vrec, 'artifact_type', ''))
            row[1].text = str(getattr(vrec, 'platform', ''))
            row[2].text = os.path.basename(str(getattr(vrec, 'source_path', '')))
            dur = getattr(vrec, 'duration_seconds', None)
            row[3].text = f"{dur:.0f} seconds" if dur else "N/A"
            row[4].text = str(getattr(vrec, 'confidence', ''))
            for cell in row:
                _style_table_cell(cell)
        _add_body(doc, "")

    # ══════════════════════════════════════════════════════════════
    # 10. SHARED LINKS (if any)
    # ══════════════════════════════════════════════════════════════
    shared_links = getattr(report, 'shared_links', [])
    if shared_links:
        _add_heading_styled(doc, "Shared AI Conversation Links")
        _add_body(doc, (
            f"The examination found {len(shared_links)} web link(s) that "
            f"point to shared AI conversations. These are links that allow "
            f"others to view an AI chat session:"
        ))
        for sl in shared_links:
            url = getattr(sl, 'url', 'N/A')
            plat = getattr(sl, 'platform', 'Unknown')
            _add_bullet(doc, f"{plat}: {url}")

    # ══════════════════════════════════════════════════════════════
    # 11. AI-GENERATED FILES (if any)
    # ══════════════════════════════════════════════════════════════
    gen_assets = getattr(report, 'generated_assets', [])
    if gen_assets:
        _add_heading_styled(doc, "Files Created by AI Services")
        _add_body(doc, (
            f"The examination identified {len(gen_assets)} file(s) that "
            f"appear to have been created or generated by an AI service:"
        ))
        for ga in gen_assets:
            path = os.path.basename(str(getattr(ga, 'asset_path', 'N/A')))
            plat = getattr(ga, 'platform', 'Unknown')
            c2pa = getattr(ga, 'c2pa_detected', False)
            c2pa_text = " (contains digital watermark confirming AI origin)" if c2pa else ""
            _add_bullet(doc, f"{path} — created by {plat}{c2pa_text}")

    # ══════════════════════════════════════════════════════════════
    # CONCLUSION
    # ══════════════════════════════════════════════════════════════
    _add_heading_styled(doc, "Conclusion")

    if report.ai_footprints:
        _add_body(doc, (
            "Based on my examination of the digital evidence, I have "
            "reached the following conclusions:"
        ))
        for fprt in report.ai_footprints:
            pname = fprt.platform.value if hasattr(fprt.platform, "value") else str(fprt.platform)
            conf = fprt.overall_confidence.value if hasattr(fprt.overall_confidence, "value") else str(fprt.overall_confidence)
            _add_body(doc, (
                f"There is evidence, with {conf} confidence, that the AI "
                f"service \"{pname}\" was accessed from the device under "
                f"examination. This conclusion is supported by "
                f"{fprt.total_artifacts} piece(s) of digital evidence, of "
                f"which {fprt.direct_artifacts} are direct evidence and "
                f"{fprt.inferred_artifacts} are supporting evidence."
            ))
    else:
        _add_body(doc, (
            "Based on my examination of the digital evidence, I did not "
            "find direct evidence that any AI service (including ChatGPT, "
            "Claude, Google Gemini, Perplexity, Copilot, Meta AI, Grok, "
            "or Poe) was used on the device under examination."
        ))
        _add_body(doc, (
            "However, I wish to make clear to the Court that the absence "
            "of evidence should not be interpreted as proof that AI services "
            "were never used. Digital evidence can be deleted, overwritten, "
            "or accessed through means (such as a different device or "
            "private browsing) that would not leave traces on this particular "
            "device."
        ))

    _add_body(doc, "")
    _add_body(doc, (
        "This report is submitted to the best of my knowledge and belief, "
        "and the analysis was conducted using accepted forensic methods and "
        "tools."
    ))

    _add_body(doc, "")
    _add_body(doc, (
        "Recommended Notice: Portions of this report were prepared with AI assistance. "
        "Review and validation by a qualified human investigator are strongly recommended "
        "before submission for any legal, forensic, or official purpose."
    ))

    # ── Examiner signature block ──────────────────────────────────
    _add_blank(doc, 2)
    _add_horizontal_line(doc)
    _add_body(doc, "Respectfully submitted,")
    _add_blank(doc, 2)
    _add_body(doc, f"____________________________________")
    _add_body(doc, examiner_display, spacing_after=Pt(2))
    _add_body(doc, "Digital Forensics Examiner", spacing_after=Pt(2))
    _add_body(doc, organization, spacing_after=Pt(2))
    _add_body(doc, f"Date: {report_date}")

    # ══════════════════════════════════════════════════════════════
    # APPENDIX — GLOSSARY OF TERMS
    # ══════════════════════════════════════════════════════════════
    _add_page_break(doc)
    _add_heading_styled(doc, "Appendix — Glossary of Terms")

    _add_body(doc, (
        "The following definitions are provided to assist non-technical "
        "readers in understanding the terms used in this report."
    ))

    glossary = [
        ("Artifact",
         "A piece of digital evidence found on a computer or device, "
         "such as a website visit record, a saved file, or a cookie. "
         "Each artifact is a clue about what happened on the device."),
        ("AI Service / AI Platform",
         "An online service that uses artificial intelligence, such as "
         "ChatGPT (by OpenAI), Claude (by Anthropic), or Gemini (by Google). "
         "These services allow a person to have conversations with a "
         "computer program that can answer questions and generate text."),
        ("Browser History",
         "A record kept by a web browser (such as Chrome, Edge, or Safari) "
         "of every website a user visits. This is one of the most direct "
         "forms of evidence that an AI service was accessed."),
        ("Cookie",
         "A small piece of data that a website stores on a user's computer. "
         "Cookies can show that a user visited or logged into a particular "
         "website, including AI services."),
        ("Confidence Level",
         "How certain the examiner is about a finding. \"High\" means the "
         "evidence strongly supports the conclusion. \"Medium\" means the "
         "evidence is suggestive but not conclusive. \"Low\" means the "
         "evidence provides only a weak indication."),
        ("Direct Evidence",
         "Evidence that clearly and unmistakably shows something happened, "
         "such as a browser history entry showing a visit to chat.openai.com."),
        ("Supporting / Corroborating Evidence",
         "Evidence that, while not proof on its own, adds strength to other "
         "findings. For example, a cookie from an AI service supports a "
         "browser history entry showing a visit to that service."),
        ("File System",
         "The way files and folders are organized on a computer's storage. "
         "Think of it as the filing cabinet system that keeps everything "
         "in order on a hard drive."),
        ("Forensic Image",
         "An exact digital copy of a computer's entire storage, including "
         "deleted files. This is the standard way evidence is preserved "
         "in digital forensic examinations."),
        ("Hash Value",
         "A digital fingerprint of a file or piece of data. If two copies "
         "of a file produce the same hash value, they are identical. This "
         "is used to prove that evidence has not been tampered with."),
        ("FRAUE",
         "Forensic Reasoning and Analysis Unit of Evidence — a single, "
         "complete finding in this report. Each FRAUE represents one "
         "instance or session of AI service usage identified by the examiner."),
        ("TRACE-AI-FR",
         "The forensic analysis tool used in this examination. It stands for "
         "\"Transparent Reporting of AI-related Claims in Evidence: A "
         "Forensic Reasoning Framework.\""),
        ("E01 (Expert Witness Format)",
         "A standard file format used to store forensic images of computer "
         "storage. It is widely accepted in courts and by law enforcement."),
        ("NTFS",
         "New Technology File System — the standard file system used by "
         "Windows computers to organize files and folders."),
        ("Byte",
         "A unit of digital storage. Computer file sizes are measured in "
         "bytes, kilobytes (KB), megabytes (MB), gigabytes (GB), or "
         "terabytes (TB)."),
    ]
    for term, definition in glossary:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run_term = p.add_run(f"{term}: ")
        _set_run(run_term, bold=True, size=_BODY_SIZE)
        run_def = p.add_run(definition)
        _set_run(run_def, size=_BODY_SIZE)

    # v5.0 Artifact Coverage Ledger Appendix
    acl = getattr(report, "artifact_coverage_ledger", [])
    gap = getattr(report, "coverage_gap_ledger", [])
    fail = getattr(report, "parse_failure_ledger", [])
    unsup = getattr(report, "unsupported_artifact_ledger", [])
    if acl or gap or fail or unsup:
        _add_page_break(doc)
        _add_heading_styled(doc, "Appendix — Artifact Coverage Ledger (v5.0)")
        if acl:
            _add_subheading(doc, "Artifact Coverage Records")
            table = doc.add_table(rows=1, cols=10)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            headers = ["Platform", "OS", "User", "Artifact Family", "Parser", "Path", "Evidence Count", "Confidence", "Reason", "Caveat"]
            for i, h in enumerate(headers):
                hdr[i].text = h
                _style_table_header(hdr[i])
            for r in acl:
                row = table.add_row().cells
                row[0].text = str(getattr(r, 'platform', ''))
                row[1].text = str(getattr(r, 'os', ''))
                row[2].text = str(getattr(r, 'user_profile', ''))
                row[3].text = str(getattr(r, 'artifact_family', ''))
                row[4].text = str(getattr(r, 'parser_used', ''))
                row[5].text = str(getattr(r, 'actual_path', ''))
                row[6].text = str(getattr(r, 'evidence_count', 0))
                row[7].text = str(getattr(r, 'confidence_impact', ''))
                row[8].text = str(getattr(r, 'reason_code', ''))
                row[9].text = str(getattr(r, 'caveat_text', ''))
                for cell in row:
                    _style_table_cell(cell)
            _add_body(doc, "")
        if gap:
            _add_subheading(doc, "Coverage Gaps")
            table = doc.add_table(rows=1, cols=8)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            headers = ["Platform", "OS", "User", "Artifact Family", "Expected Path", "Reason", "Confidence", "Caveat"]
            for i, h in enumerate(headers):
                hdr[i].text = h
                _style_table_header(hdr[i])
            for r in gap:
                row = table.add_row().cells
                row[0].text = str(getattr(r, 'platform', ''))
                row[1].text = str(getattr(r, 'os', ''))
                row[2].text = str(getattr(r, 'user_profile', ''))
                row[3].text = str(getattr(r, 'artifact_family', ''))
                row[4].text = str(getattr(r, 'expected_path', ''))
                row[5].text = str(getattr(r, 'gap_reason', ''))
                row[6].text = str(getattr(r, 'confidence_impact', ''))
                row[7].text = str(getattr(r, 'caveat_text', ''))
                for cell in row:
                    _style_table_cell(cell)
            _add_body(doc, "")
        if fail:
            _add_subheading(doc, "Parse Failures")
            table = doc.add_table(rows=1, cols=8)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            headers = ["Platform", "OS", "User", "Artifact Family", "Path", "Parser", "Error", "Caveat"]
            for i, h in enumerate(headers):
                hdr[i].text = h
                _style_table_header(hdr[i])
            for r in fail:
                row = table.add_row().cells
                row[0].text = str(getattr(r, 'platform', ''))
                row[1].text = str(getattr(r, 'os', ''))
                row[2].text = str(getattr(r, 'user_profile', ''))
                row[3].text = str(getattr(r, 'artifact_family', ''))
                row[4].text = str(getattr(r, 'path', ''))
                row[5].text = str(getattr(r, 'parser_used', ''))
                row[6].text = str(getattr(r, 'error_message', ''))
                row[7].text = str(getattr(r, 'caveat_text', ''))
                for cell in row:
                    _style_table_cell(cell)
            _add_body(doc, "")
        if unsup:
            _add_subheading(doc, "Unsupported Artifact Locations")
            table = doc.add_table(rows=1, cols=7)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            headers = ["Platform", "OS", "User", "Artifact Family", "Path", "Parser Needed", "Notes"]
            for i, h in enumerate(headers):
                hdr[i].text = h
                _style_table_header(hdr[i])
            for r in unsup:
                row = table.add_row().cells
                row[0].text = str(getattr(r, 'platform', ''))
                row[1].text = str(getattr(r, 'os', ''))
                row[2].text = str(getattr(r, 'user_profile', ''))
                row[3].text = str(getattr(r, 'artifact_family', ''))
                row[4].text = str(getattr(r, 'path', ''))
                row[5].text = str(getattr(r, 'parser_needed', ''))
                row[6].text = str(getattr(r, 'notes', ''))
                for cell in row:
                    _style_table_cell(cell)
            _add_body(doc, "")
    # ── AI Tool Inventory Checklist (v5.0) ────────────────────────
    checklist = getattr(report, "forensic_checklist", [])
    if checklist:
        _add_page_break(doc)
        _add_heading_styled(doc, "AI Tool Inventory Checklist")

        _add_body(doc, (
            "As part of the forensic examination, the examiner scanned the "
            "evidence for signs of all known AI tools — including desktop "
            "applications, IDE extensions, local Large Language Models (LLMs), "
            "terminal CLI tools, and browser-accessed AI services."
        ))

        _add_body(doc, (
            "IMPORTANT: The mere presence of an AI tool installation, browser "
            "cache entry, or configuration file does not prove that the tool "
            "was actually used for any particular purpose. Installation is not "
            "the same as use. Furthermore, tool presence alone cannot identify "
            "which person used the tool."
        ))

        n_found = sum(
            1 for e in checklist
            if str(getattr(e, "evidence_status", "")).replace("EvidenceStatus.", "") == "FOUND"
        )
        _add_body(doc, (
            f"Of {len(checklist)} registered AI tools in the forensic inventory, "
            f"{n_found} tool(s) were detected in the evidence."
        ))

        # Summary table
        _add_subheading(doc, "Tool Detection Summary")
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Tool Name", "Category", "Status", "Confidence", "Key Caveats"]):
            hdr[i].text = h
            _style_table_header(hdr[i])

        for entry in checklist:
            row = table.add_row().cells
            row[0].text = str(getattr(entry, "tool_name", ""))
            row[1].text = str(getattr(entry, "category", ""))
            status = str(getattr(entry, "evidence_status", "")).replace("EvidenceStatus.", "")
            row[2].text = status
            row[3].text = str(getattr(entry, "confidence", "N/A") or "N/A")
            caveats = [str(f).replace("CaveatFlag.", "") for f in (getattr(entry, "caveat_flags", []) or [])]
            row[4].text = ", ".join(caveats[:3]) if caveats else "—"
            for cell in row:
                _style_table_cell(cell)

        _add_body(doc, "")

        # Detail for FOUND tools
        found = [
            e for e in checklist
            if str(getattr(e, "evidence_status", "")).replace("EvidenceStatus.", "") == "FOUND"
        ]
        if found:
            _add_subheading(doc, "Detected Tools — Detail")
            for entry in found:
                _add_body(doc, (
                    f"{entry.tool_name} ({entry.category}) — "
                    f"detected via {getattr(entry, 'detection_method', '—') or '—'}. "
                    f"Execution surface: {getattr(entry, 'execution_surface', '—') or '—'}. "
                    f"Artifact count: {getattr(entry, 'artifact_count', 0)}."
                ))
                if getattr(entry, "artifact_paths", None):
                    for p in entry.artifact_paths[:5]:
                        _add_bullet(doc, f"Path: {p}")
                caveats = [str(f).replace("CaveatFlag.", "") for f in (getattr(entry, "caveat_flags", []) or [])]
                if caveats:
                    _add_bullet(doc, f"Caveats: {', '.join(caveats)}")
                # Add resolved caveat language
                try:
                    from .caveats import get_caveat_text
                    caveat_texts = get_caveat_text(getattr(entry, "caveat_flags", []) or [])
                    for ct in caveat_texts:
                        _add_bullet(doc, f"⚠ {ct}")
                except Exception:
                    pass

        # Global evidentiary caveat footer
        try:
            from .caveats import GLOBAL_REPORT_FOOTER
            _add_body(doc, "")
            _add_body(doc, f"Evidentiary Notice: {GLOBAL_REPORT_FOOTER}")
        except Exception:
            pass

    # ── Save ──────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    logger.info("Word report saved to %s", output_path)
    return output_path


def _human_size(nbytes: int) -> str:
    """Format bytes as human-readable string."""
    for unit in ("B", "KB", "MB", "GB"):
        if nbytes < 1024:
            return f"{nbytes:.0f} {unit}" if unit == "B" else f"{nbytes:.1f} {unit}"
        nbytes /= 1024
    return f"{nbytes:.1f} TB"
